"""Word/PDF export helpers for exam papers with unit structure.

V3.5.1: Loads questions from ExamPaperUnit + ExamPaperUnitQuestion instead
of the deleted exam_paper_questions association table.
"""
import io
import json
import re
from urllib.parse import quote
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from app.models.exam_paper import ExamPaper, ExamPaperUnit, ExamPaperUnitQuestion
from docx.shared import Pt, Cm


TYPE_LABELS = {
    "FILL_BLANK": "填空题",
    "SINGLE_CHOICE": "单选题",
    "MULTIPLE_CHOICE": "多选题",
    "SUBJECTIVE": "解答题",
}


def _normalize_options(title: str, options: list | None) -> tuple[str, list[dict] | None]:
    """规范化选项：纯字母选项补文本，题干裁剪行内选项。返回 (title, options)。"""
    if not options:
        return title, None

    # 纯字母选项 ["A","B","C","D"] → 从题干提取文本
    if all(isinstance(o, str) and len(o) == 1 and o in 'ABCDEFGH' for o in options):
        m = re.search(r'\s*A[.．、）\)]', title)
        if m and m.start() > 0:
            stripped = title[m.start():].strip()
            title = title[:m.start()].rstrip(' （(').strip()
            parts = re.split(r'\s+(?=[A-H][.．、）\)])', stripped)
            full_opts = []
            for part in parts:
                pm = re.match(r'^([A-H])[.．、）)]\s*(.*)', part)
                if pm:
                    full_opts.append({'label': pm.group(1), 'text': pm.group(2).strip()})
            if len(full_opts) == len(options):
                return title, full_opts
        return title, options

    # 字符串选项 "A. text" → 对象 {label,text}
    if all(isinstance(o, str) for o in options):
        full_opts = []
        for o in options:
            pm = re.match(r'^([A-H])[.．、）)]\s*(.*)', o)
            if pm:
                full_opts.append({'label': pm.group(1), 'text': pm.group(2)})
            else:
                full_opts.append({'label': '', 'text': o})
        # 同时裁剪题干
        m = re.search(r'\s*A[.．、）\)]', title)
        if m and m.start() > 0:
            title = title[:m.start()].rstrip(' （(').strip()
        return title, full_opts

    return title, options


def _parse_question(question, uq_score: int) -> dict:
    """Parse a single Question row into the export dict format."""
    correct_answer = question.correct_answer or ""
    options = None
    answer_text = ""
    try:
        ad = json.loads(correct_answer)
        if isinstance(ad, dict):
            options = ad.get("options")
            answer_text = ad.get("correct_answer", "")
            if isinstance(answer_text, list):
                answer_text = ", ".join(str(x) for x in answer_text)
            else:
                answer_text = str(answer_text)
    except (json.JSONDecodeError, TypeError):
        answer_text = correct_answer

    title, options = _normalize_options(question.title or "", options)

    return {
        "title": title,
        "question_type": question.question_type,
        "type_label": TYPE_LABELS.get(question.question_type, question.question_type),
        "difficulty": question.difficulty or "",
        "score": uq_score or question.score or 0,
        "correct_answer": correct_answer,
        "answer_text": answer_text,
        "options": options,
        "explanation": question.explanation or "",
    }


async def load_paper_with_questions(
    exam_paper_id,
    db: AsyncSession,
):
    """Load a paper together with all units and their questions.

    Returns (paper, questions) where questions is a flat list of dicts with
    a sequential ``index``.
    """
    result = await db.execute(
        select(ExamPaper)
        .where(ExamPaper.id == exam_paper_id)
        .options(
            selectinload(ExamPaper.units)
            .selectinload(ExamPaperUnit.questions)
            .selectinload(ExamPaperUnitQuestion.question),
        )
    )
    paper = result.scalar_one_or_none()
    if not paper:
        from fastapi import HTTPException, status
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Exam paper not found",
        )

    questions = []
    for unit in paper.units:
        for uq in unit.questions:
            parsed = _parse_question(uq.question, uq.score)
            parsed["index"] = len(questions) + 1
            parsed["unit_name"] = unit.name
            questions.append(parsed)

    return paper, questions


# 字体字号：二号22pt 四号14pt 小四12pt 五号10.5pt 小五9pt
FONT_TITLE_SIZE = Pt(22)
FONT_SECTION_SIZE = Pt(14)
FONT_QNUM_SIZE = Pt(12)
FONT_BODY_SIZE = Pt(10.5)
FONT_SMALL_SIZE = Pt(9)

def _set_cn_font(run, cn_name="宋体", en_name="Times New Roman"):
    """设置中西文字体"""
    from docx.oxml.ns import qn
    run.font.name = en_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_name)

def _add_divider(doc):
    """1磅黑色实线分割线"""
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '8')
    bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def _write_type_sections(doc, qs, type_order, Cm, Pt, RGBColor):
    """写入题型分区（Word）"""
    grouped: dict[str, list] = {}
    for q in qs:
        grouped.setdefault(q["question_type"], []).append(q)
    num_labels = ['一', '二', '三', '四', '五', '六', '七', '八']
    section_idx = 0
    for t in type_order:
        tqs = grouped.get(t, [])
        if not tqs:
            continue
        header = doc.add_paragraph()
        type_score = sum(q['score'] for q in tqs)
        per_q = tqs[0]['score'] if tqs else 0
        num = num_labels[section_idx] if section_idx < len(num_labels) else str(section_idx + 1)
        section_idx += 1
        # 模块标题 — 黑体四号左对齐加粗
        header_text = f"{num}、{tqs[0]['type_label']}"
        # 答题要求（选择题）
        if t == "SINGLE_CHOICE":
            header_text += f"（每题{per_q}分，共{len(tqs)}题，合计{type_score}分。在每小题给出的选项中，只有一项符合题目要求）"
        elif t == "MULTIPLE_CHOICE":
            header_text += f"（每题{per_q}分，共{len(tqs)}题，合计{type_score}分。在每小题给出的选项中，有多项符合题目要求）"
        elif t == "FILL_BLANK":
            header_text += f"（每题{per_q}分，共{len(tqs)}题，合计{type_score}分）"
        elif t == "SUBJECTIVE":
            header_text += f"（共{len(tqs)}题，合计{type_score}分）"
        run = header.add_run(header_text)
        run.bold = True
        run.font.size = FONT_SECTION_SIZE
        _set_cn_font(run, "黑体")
        # 1磅黑色实线分割
        from docx.oxml import OxmlElement
        pPr = header._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
        bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '8')
        bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        for q in tqs:
            _write_question_word(doc, q, t, Cm, Pt)
        # 题型间加倍行距
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(12)


def _write_question_word(doc, q, t, Cm, Pt):
    """写入单道题目（Word）— 高考标准排版"""
    is_choice = t in ("SINGLE_CHOICE", "MULTIPLE_CHOICE")

    # 题号+题干 — 黑体小四号加粗题号, 宋体五号题干, 1.5倍行距
    q_para = doc.add_paragraph()
    q_para.paragraph_format.space_after = Pt(0)
    q_para.paragraph_format.line_spacing = 1.5
    q_para.paragraph_format.keep_with_next = True  # 与选项保持同页
    # 选择题用顿号"1、"，非选择题用句号"6."
    num_str = f"{q['index']}、"
    run_num = q_para.add_run(num_str)
    run_num.bold = True
    run_num.font.size = FONT_QNUM_SIZE
    _set_cn_font(run_num, "黑体")
    # 题干
    title_text = f"{q['title']} "
    run_title = q_para.add_run(title_text)
    run_title.font.size = FONT_BODY_SIZE
    _set_cn_font(run_title)
    # 分值 — 右侧小五号
    run_score = q_para.add_run(f"（{q['score']}分）")
    run_score.font.size = FONT_SMALL_SIZE
    _set_cn_font(run_score)

    if is_choice and q["options"] and len(q["options"]) > 0:
        for i, opt in enumerate(q["options"]):
            opt_para = doc.add_paragraph()
            opt_para.paragraph_format.left_indent = Cm(0.74)
            opt_para.paragraph_format.space_after = Pt(0)
            opt_para.paragraph_format.space_before = Pt(0)
            opt_para.paragraph_format.line_spacing = 1.5
            if i < len(q["options"]) - 1:
                opt_para.paragraph_format.keep_with_next = True  # 选项间不分页
            if isinstance(opt, dict):
                label = opt.get('label', '')
                text = opt.get('text', '')
                run_ol = opt_para.add_run(f"{label}、")
                run_ol.bold = True
                run_ol.font.size = FONT_BODY_SIZE
                _set_cn_font(run_ol)
                run_ot = opt_para.add_run(text)
                run_ot.font.size = FONT_BODY_SIZE
                _set_cn_font(run_ot)
            else:
                line = str(opt)
                pm = re.match(r'^([A-H])[.．、）)]\s*(.*)', line)
                label = pm.group(1) if pm else ''
                text = pm.group(2) if pm else line
                run_ol = opt_para.add_run(f"{label}、")
                run_ol.bold = True
                run_ol.font.size = FONT_BODY_SIZE
                _set_cn_font(run_ol)
                run_ot = opt_para.add_run(text)
                run_ot.font.size = FONT_BODY_SIZE
                _set_cn_font(run_ot)
    elif t == "SUBJECTIVE":
        doc.add_paragraph()  # 页首留白
        for _ in range(13):
            space_para = doc.add_paragraph()
            space_para.paragraph_format.line_spacing = 1.5
            space_para.add_run(" ").font.size = FONT_BODY_SIZE
    elif t == "FILL_BLANK":
        doc.add_paragraph().paragraph_format.line_spacing = 1.5


async def export_word(exam_paper_id, db: AsyncSession):
    """Generate a Word document for the exam paper."""
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paper, questions = await load_paper_with_questions(exam_paper_id, db)

    doc = Document()
    # ── A4 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)  # 20mm + 10mm装订线
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(2.0)  # 页眉与正文间距

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = FONT_BODY_SIZE
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.5
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 页眉 ──
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run(paper.title or "")
    hr.font.size = FONT_SMALL_SIZE
    _set_cn_font(hr, "宋体")

    # ── 页脚 ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from docx.oxml import OxmlElement
    # 页码域代码
    run_page = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
    run_page._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set('{http://schemas.openxmlformats.org/word}space', 'preserve')
    instrText.text = ' PAGE '
    run_page._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
    run_page._r.append(fldChar2)
    fp.add_run(" / ").font.size = FONT_SMALL_SIZE
    run_total = fp.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
    run_total._r.append(fldChar3)
    instrText2 = OxmlElement('w:instrText')
    instrText2.set('{http://schemas.openxmlformats.org/word}space', 'preserve')
    instrText2.text = ' NUMPAGES '
    run_total._r.append(instrText2)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
    run_total._r.append(fldChar4)

    # 大标题 — 黑体二号居中加粗
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    run_t = title_para.add_run(paper.title or "")
    run_t.bold = True
    run_t.font.size = FONT_TITLE_SIZE
    _set_cn_font(run_t, "黑体")

    if paper.subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.paragraph_format.space_after = Pt(2)
        run_s = sub_para.add_run(paper.subtitle)
        run_s.font.size = FONT_SECTION_SIZE
        _set_cn_font(run_s)

    # 信息行
    info_parts = []
    if paper.subject: info_parts.append(paper.subject)
    if paper.grade_level and isinstance(paper.grade_level, dict):
        grades = paper.grade_level.get("grades", [])
        if grades: info_parts.append(', '.join(grades))
    info_parts.append(f"总分: {paper.total_score}分")
    if paper.duration_minutes: info_parts.append(f"时长: {paper.duration_minutes}分钟")
    if info_parts:
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.paragraph_format.space_after = Pt(4)
        run_i = info_para.add_run(" | ".join([p for p in info_parts if p]))
        run_i.font.size = FONT_SMALL_SIZE
        _set_cn_font(run_i)

    # 注意事项
    if paper.instructions:
        note_para = doc.add_paragraph()
        note_para.paragraph_format.space_after = Pt(4)
        run_n = note_para.add_run(f"注意事项：{paper.instructions}")
        run_n.font.size = FONT_SMALL_SIZE
        _set_cn_font(run_n)

    _add_divider(doc)

    # ── 题型内容 ──
    type_order = ["FILL_BLANK", "SINGLE_CHOICE", "MULTIPLE_CHOICE", "SUBJECTIVE"]

    if getattr(paper, 'show_units', False):
        unit_groups: dict[str, list] = {}
        for q in questions:
            un = q.get("unit_name", "")
            unit_groups.setdefault(un, []).append(q)
        for uname, uqs in unit_groups.items():
            u_header = doc.add_paragraph()
            u_header.paragraph_format.space_before = Pt(6)
            u_score = sum(q['score'] for q in uqs)
            run_u = u_header.add_run(f"{uname}（共{len(uqs)}题，{u_score}分）")
            run_u.bold = True
            run_u.font.size = FONT_SECTION_SIZE
            _set_cn_font(run_u, "黑体")
            _write_type_sections(doc, uqs, type_order, Cm, Pt, RGBColor)
    else:
        _write_type_sections(doc, questions, type_order, Cm, Pt, RGBColor)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = quote((paper.title or "exam_paper") + ".docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=utf-8''{safe_name}"
        },
    )


def _write_q_pdf(pdf, q, t):
    """写入单道题目（PDF）"""
    is_choice = t in ("SINGLE_CHOICE", "MULTIPLE_CHOICE")
    # 估空间：选择题(选项×7mm+20)，解答题(80mm)，填空(15mm)
    if is_choice:
        needed = len(q.get("options", [])) * 7 + 20
    elif t == "SUBJECTIVE":
        needed = 80
    else:
        needed = 15
    if pdf.get_y() + needed > pdf.h - pdf.b_margin:
        pdf.add_page()
    num_str = f"{q['index']}、"
    pdf.set_font("CJK", "B", 12)
    pdf.cell(8, 7, num_str)
    pdf.set_font("CJK", "", 10.5)
    title_line = f"{q['title']}（{q['score']}分）"
    pdf.multi_cell(0, 7, title_line)
    pdf.ln(1)
    if is_choice and q["options"] and len(q["options"]) > 0:
        # 估空间：选项数×7mm行高，不够就换页
        needed = len(q["options"]) * 7 + 5
        if pdf.get_y() + needed > pdf.h - pdf.b_margin:
            pdf.add_page()
        pdf.set_font("CJK", "", 10.5)
        for opt in q["options"]:
            if isinstance(opt, dict):
                label = opt.get('label', '')
                text = opt.get('text', '')
                pdf.cell(8, 6, "")
                pdf.cell(0, 6, f"{label}、{text}", new_x="LMARGIN", new_y="NEXT")
            else:
                line = str(opt)
                pm = re.match(r'^([A-H])[.．、）)]\s*(.*)', line)
                label = pm.group(1) if pm else ''
                text = pm.group(2) if pm else line
                pdf.cell(8, 6, "")
                pdf.cell(0, 6, f"{label}、{text}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)
    elif t == "SUBJECTIVE":
        pdf.ln(70)
    elif t == "FILL_BLANK":
        pdf.ln(4)

def _write_type_sections_pdf(pdf, qs, type_order):
    """写入题型分区（PDF）"""
    grouped = {}
    for q in qs:
        grouped.setdefault(q["question_type"], []).append(q)
    num_labels = ['一', '二', '三', '四', '五', '六', '七', '八']
    section_idx = 0
    for t in type_order:
        tqs = grouped.get(t, [])
        if not tqs:
            continue
        type_score = sum(q['score'] for q in tqs)
        per_q = tqs[0]['score'] if tqs else 0
        num = num_labels[section_idx] if section_idx < len(num_labels) else str(section_idx + 1)
        section_idx += 1
        pdf.set_font("CJK", "B", 14)
        header_text = f"{num}、{tqs[0]['type_label']}（每题{per_q}分，共{len(tqs)}题，合计{type_score}分）"
        pdf.cell(0, 10, header_text, new_x="LMARGIN", new_y="NEXT")
        y = pdf.get_y()
        pdf.set_draw_color(0)
        pdf.set_line_width(0.1)
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.ln(3)
        for q in tqs:
            _write_q_pdf(pdf, q, t)
        pdf.ln(6)

async def export_pdf(exam_paper_id, db: AsyncSession):
    """Generate a PDF for the exam paper."""
    from fpdf import FPDF

    paper, questions = await load_paper_with_questions(exam_paper_id, db)

    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.set_auto_page_break(True, 25)
    pdf.add_page()
    pdf.set_left_margin(30)
    pdf.set_right_margin(20)
    pdf.set_top_margin(30)

    CJK_PATHS = [
        "/usr/share/fonts/truetype/arphic-gbsn00lp/gbsn00lp.ttf",
        "/usr/share/fonts/truetype/arphic-gkai00mp/gkai00mp.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    font_loaded = False
    for fp in CJK_PATHS:
        try:
            pdf.add_font("CJK", "", fp)
            pdf.add_font("CJK", "B", fp)
            font_loaded = True
            break
        except Exception:
            continue
    if not font_loaded:
        pdf.set_font("Helvetica", "", 12)

    def header():
        pdf.set_font("CJK", "", 9)
        pdf.cell(0, 8, paper.title or "", align="R")
        pdf.ln(8)
    pdf.header = header

    def footer():
        pdf.set_y(-20)
        pdf.set_font("CJK", "", 9)
        pdf.cell(0, 10, f"第{pdf.page_no()}页", align="C")
    pdf.footer = footer

    pdf.set_font("CJK", "B", 22)
    pdf.cell(0, 14, paper.title or "", new_x="LMARGIN", new_y="NEXT", align="C")
    if paper.subtitle:
        pdf.set_font("CJK", "", 14)
        pdf.cell(0, 10, paper.subtitle, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(2)

    pdf.set_font("CJK", "", 9)
    info_parts = []
    if paper.subject: info_parts.append(paper.subject)
    if paper.grade_level and isinstance(paper.grade_level, dict):
        grades = paper.grade_level.get("grades", [])
        if grades: info_parts.append(', '.join(grades))
    info_parts.append(f"总分: {paper.total_score}分")
    if paper.duration_minutes: info_parts.append(f"时长: {paper.duration_minutes}分钟")
    if info_parts:
        pdf.cell(0, 8, " | ".join([p for p in info_parts if p]), new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(2)

    pdf.ln(4)

    type_order = ["FILL_BLANK", "SINGLE_CHOICE", "MULTIPLE_CHOICE", "SUBJECTIVE"]
    if getattr(paper, 'show_units', False):
        unit_groups = {}
        for q in questions:
            unit_groups.setdefault(q.get("unit_name", ""), []).append(q)
        for uname, uqs in unit_groups.items():
            pdf.set_font("CJK", "B", 14)
            u_score = sum(q['score'] for q in uqs)
            pdf.cell(0, 10, f"{uname}（共{len(uqs)}题，{u_score}分）", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
            _write_type_sections_pdf(pdf, uqs, type_order)
    else:
        _write_type_sections_pdf(pdf, questions, type_order)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    safe_name = quote((paper.title or "exam_paper") + ".pdf")
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{safe_name}"},
    )
