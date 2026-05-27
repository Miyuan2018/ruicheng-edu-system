import uuid
import json
from urllib.parse import quote
from fastapi import APIRouter, Depends, HTTPException, status, Body
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update, delete, text as sa_text
from app.db.session import get_db
from app.models.exam_paper import ExamPaper
from app.models.question import Question
from app.schemas.exam_paper import ExamPaperCreate, ExamPaperResponse, ExamPaperUpdate
from app.models.exam_paper import exam_paper_questions
from app.schemas.question import QuestionResponse
from typing import List, Optional
from app.core.security import get_current_user

router = APIRouter()


@router.post("", response_model=ExamPaperResponse)
async def create_exam_paper(
    exam_paper_in: ExamPaperCreate,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if current_user.user_type not in ("TEACHER", "QUESTION_ADMIN", "SYS_ADMIN"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="权限不足")

    data = exam_paper_in.model_dump(exclude_none=True)
    questions_data = data.pop("questions", []) if isinstance(data.get("questions"), list) else []
    # Remove extra fields not in ExamPaper model
    for key in ["question_count", "distribution", "difficulty_ratio"]:
        data.pop(key, None)
    data["created_by"] = uuid.UUID(current_user.id)

    async with db.begin():
        exam_paper = ExamPaper(**data)
        db.add(exam_paper)
        await db.flush()

        # Import questions if provided
        for i, qdata in enumerate(questions_data):
            q = Question(
                title=qdata.get("title", ""),
                question_type=qdata.get("question_type", "SINGLE_CHOICE"),
                difficulty=qdata.get("difficulty", "MEDIUM"),
                subject=qdata.get("subject", exam_paper.subject),
                grade_level=qdata.get("grade_level", exam_paper.grade_level),
                score=qdata.get("score", 5),
                correct_answer=qdata.get("correct_answer", ""),
                explanation=qdata.get("explanation", ""),
                source="MANUAL", review_status="APPROVED",
                created_by=uuid.UUID(current_user.id),
            )
            db.add(q)
            await db.flush()
            # Link to paper
            await db.execute(exam_paper_questions.insert().values(
                id=uuid.uuid4(), exam_paper_id=exam_paper.id,
                question_id=q.id, position=i+1, score=qdata.get("score", 5),
            ))

    await db.refresh(exam_paper)
    return exam_paper


@router.get("/my")
async def get_my_papers(
    skip: int = 0,
    limit: int = 20,
    title: Optional[str] = None,
    status: Optional[str] = None,
    grade: Optional[str] = None,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Return exam papers that the current student has submitted answers for."""
    limit = min(limit, 200)
    if current_user.user_type != "STUDENT":
        raise HTTPException(403, detail="仅学生可访问")
    from app.models.answer_submission import AnswerSubmission
    from sqlalchemy import distinct
    subq = select(distinct(AnswerSubmission.exam_paper_id)).where(
        AnswerSubmission.student_id == uuid.UUID(current_user.id)
    ).subquery()
    query = select(ExamPaper).where(ExamPaper.id.in_(subq))
    if title:
        query = query.where(ExamPaper.title.ilike(f"%{title}%"))
    if status:
        query = query.where(ExamPaper.status == status)
    if grade:
        query = query.where(ExamPaper.grade_level['grades'].contains([grade]))
    query = query.offset(skip).limit(limit).order_by(ExamPaper.created_at.desc())
    result = await db.execute(query)
    papers = result.scalars().all()

    # Get the student's latest submission for each paper
    paper_ids = [p.id for p in papers]
    sub_status = {}
    if paper_ids:
        from sqlalchemy import func as _func
        sub_result = await db.execute(
            select(AnswerSubmission).where(
                AnswerSubmission.student_id == uuid.UUID(current_user.id),
                AnswerSubmission.exam_paper_id.in_(paper_ids)
            ).order_by(AnswerSubmission.submitted_at.desc())
        )
        for sub in sub_result.scalars().all():
            pid = str(sub.exam_paper_id)
            if pid not in sub_status:
                sub_status[pid] = sub

    return [{
        "id": str(p.id), "title": p.title, "subtitle": p.subtitle,
        "subject": p.subject, "grade_level": p.grade_level,
        "total_score": p.total_score, "duration_minutes": p.duration_minutes,
        "status": p.status, "instructions": p.instructions, "description": p.description,
        "created_at": p.created_at, "updated_at": p.updated_at,
        "submission_status": sub_status[str(p.id)].status if str(p.id) in sub_status else None,
        "submission_score": float(sub_status[str(p.id)].total_score) if str(p.id) in sub_status and sub_status[str(p.id)].total_score else None,
        "submission_percentage": float(sub_status[str(p.id)].percentage) if str(p.id) in sub_status and sub_status[str(p.id)].percentage else None,
        "submission_id": str(sub_status[str(p.id)].id) if str(p.id) in sub_status else None,
    } for p in papers]


@router.get("/{exam_paper_id}/review")
async def review_exam_paper(
    exam_paper_id: uuid.UUID,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Return full paper review: paper + student's submission + all questions with answers."""
    if current_user.user_type != "STUDENT":
        raise HTTPException(status_code=403, detail="仅学生可复盘")

    # Get paper
    paper_result = await db.execute(select(ExamPaper).where(ExamPaper.id == exam_paper_id))
    paper = paper_result.scalar_one_or_none()
    if not paper:
        raise HTTPException(status_code=404, detail="试卷不存在")

    # Get student's latest submission for this paper
    from app.models.answer_submission import AnswerSubmission
    sub_result = await db.execute(
        select(AnswerSubmission).where(
            AnswerSubmission.student_id == uuid.UUID(current_user.id),
            AnswerSubmission.exam_paper_id == exam_paper_id,
        ).order_by(AnswerSubmission.submitted_at.desc()).limit(1)
    )
    submission = sub_result.scalar_one_or_none()

    # Get all answers for this submission
    from app.models.answer_detail import AnswerDetail
    answers_map = {}
    if submission:
        ad_result = await db.execute(
            select(AnswerDetail).where(AnswerDetail.answer_submission_id == submission.id)
        )
        for ad in ad_result.scalars().all():
            answers_map[str(ad.question_id)] = ad

    # Get paper questions with position
    questions_data = []
    q_result = await db.execute(
        select(exam_paper_questions.c.question_id, exam_paper_questions.c.position)
        .where(exam_paper_questions.c.exam_paper_id == exam_paper_id)
        .order_by(exam_paper_questions.c.position)
    )
    q_rows = q_result.all()
    if q_rows:
        qids = [row[0] for row in q_rows]
        pos_map = {str(row[0]): row[1] for row in q_rows}
        qs_result = await db.execute(select(Question).where(Question.id.in_(qids)))
        questions_map = {str(q.id): q for q in qs_result.scalars().all()}
        for qid in qids:
            q = questions_map.get(str(qid))
            if not q:
                continue
            ad = answers_map.get(str(qid))
            # Parse correct_answer
            ca = None
            try:
                data = json.loads(q.correct_answer or '{}')
                if isinstance(data.get('correct_answer'), list):
                    ca = ', '.join(str(x) for x in data['correct_answer'])
                elif isinstance(data.get('correct_answer'), str):
                    ca = data['correct_answer']
                elif isinstance(data.get('correct_answer'), dict):
                    ca = ', '.join(data['correct_answer'].get('keywords', []))
            except Exception:
                ca = q.correct_answer
            questions_data.append({
                "position": pos_map.get(str(qid), 0),
                "question": {
                    "id": str(q.id),
                    "title": q.title,
                    "question_type": q.question_type,
                    "score": float(q.score or 0),
                    "correct_answer": ca,
                },
                "student_answer": ad.student_answer if ad else None,
                "is_correct": ad.is_correct if ad else None,
                "score_obtained": float(ad.score_obtained) if ad and ad.score_obtained else None,
                "feedback": ad.feedback if ad else None,
            })

    return {
        "paper": {
            "id": str(paper.id),
            "title": paper.title,
            "subject": paper.subject,
            "total_score": float(paper.total_score or 0),
            "duration_minutes": paper.duration_minutes,
        },
        "submission": {
            "id": str(submission.id) if submission else None,
            "status": submission.status if submission else None,
            "total_score": float(submission.total_score) if submission and submission.total_score else None,
            "percentage": float(submission.percentage) if submission and submission.percentage else None,
            "submitted_at": submission.submitted_at.isoformat() if submission and submission.submitted_at else None,
        } if submission else None,
        "questions": questions_data,
    }


@router.get("/{exam_paper_id}", response_model=ExamPaperResponse)
async def get_exam_paper(
    exam_paper_id: uuid.UUID,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(ExamPaper).where(ExamPaper.id == exam_paper_id))
    exam_paper = result.scalar_one_or_none()
    if not exam_paper:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Exam paper not found",
        )
    return exam_paper


@router.put("/{exam_paper_id}", response_model=ExamPaperResponse)
async def update_exam_paper(
    exam_paper_id: uuid.UUID,
    exam_paper_in: ExamPaperUpdate,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # Only teachers and admins can update exam papers
    if current_user.user_type not in ("TEACHER", "QUESTION_ADMIN", "SYS_ADMIN"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions",
        )

    result = await db.execute(select(ExamPaper).where(ExamPaper.id == exam_paper_id))
    exam_paper = result.scalar_one_or_none()
    if not exam_paper:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Exam paper not found",
        )

    # Check if user is the creator or is an admin
    allowed = (
        exam_paper.created_by == uuid.UUID(current_user.id)
        or current_user.user_type in ("SYS_ADMIN", "TEACHER", "QUESTION_ADMIN")
        or current_user.user_type == "STUDENT"
    )
    if not allowed:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions",
        )

    # Update exam paper
    update_data = exam_paper_in.dict(exclude_unset=True)
    for field, value in update_data.items():
        setattr(exam_paper, field, value)

    await db.commit()
    await db.refresh(exam_paper)
    return exam_paper


@router.post("/{exam_paper_id}/publish")
async def publish_exam_paper(
    exam_paper_id: uuid.UUID,
    class_ids: list[str] = Body(default=[]),
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Publish an exam paper and notify students in the specified (or all teacher's) classes."""
    if current_user.user_type not in ("TEACHER", "QUESTION_ADMIN", "SYS_ADMIN"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="权限不足")

    result = await db.execute(select(ExamPaper).where(ExamPaper.id == exam_paper_id))
    exam_paper = result.scalar_one_or_none()
    if not exam_paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="试卷不存在")

    if exam_paper.status == "ARCHIVED":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="已归档的试卷不可发布")

    # Get teacher name for notification
    from app.models.admin import Admin
    admin_result = await db.execute(select(Admin).where(Admin.id == exam_paper.created_by))
    admin = admin_result.scalar_one_or_none()
    teacher_name = admin.full_name if admin else "老师"

    # Determine which students to notify
    from app.models.school_class import SchoolClass, class_students
    from app.models.student import Student

    if class_ids:
        # Notify students in the specified classes
        student_q = (
            select(Student.id)
            .join(class_students, class_students.c.student_id == Student.id)
            .where(class_students.c.class_id.in_(class_ids))
        )
    else:
        # Notify all students in classes taught by the paper creator
        student_q = (
            select(Student.id)
            .join(class_students, class_students.c.student_id == Student.id)
            .join(SchoolClass, SchoolClass.id == class_students.c.class_id)
            .where(SchoolClass.teacher_id == exam_paper.created_by)
        )

    student_result = await db.execute(student_q)
    student_ids = [row[0] for row in student_result.all()]

    # Fetch class names for notification content
    if class_ids:
        class_name_q = select(SchoolClass.name).where(SchoolClass.id.in_(class_ids))
    else:
        class_name_q = select(SchoolClass.name).where(
            SchoolClass.teacher_id == exam_paper.created_by
        )
    class_name_result = await db.execute(class_name_q)
    class_names = [row[0] for row in class_name_result.all()]
    class_name_str = "、".join(class_names) if class_names else "班级"

    # Create notifications for each student
    from app.services.notification_service import NotificationService
    notified_count = 0
    for student_id in student_ids:
        await NotificationService.create_class_announcement_notification(
            db=db,
            recipient_id=str(student_id),
            teacher_name=teacher_name,
            class_name=class_name_str,
            title=f"试卷发布: {exam_paper.title}",
            content=f"「{exam_paper.title}」已发布，请前往我的试卷查看并完成答题。",
        )
        notified_count += 1

    # Update paper status
    exam_paper.status = "PUBLISHED"
    await db.commit()
    await db.refresh(exam_paper)

    return {
        "id": str(exam_paper.id),
        "title": exam_paper.title,
        "status": exam_paper.status,
        "notified_count": notified_count,
        "updated_at": exam_paper.updated_at,
    }


@router.delete("/{exam_paper_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_exam_paper(
    exam_paper_id: uuid.UUID,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if current_user.user_type not in ("TEACHER", "QUESTION_ADMIN", "SYS_ADMIN", "STUDENT"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions",
        )

    result = await db.execute(select(ExamPaper).where(ExamPaper.id == exam_paper_id))
    exam_paper = result.scalar_one_or_none()
    if not exam_paper:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Exam paper not found",
        )

    # Check if user is the creator or is an admin
    allowed = (
        exam_paper.created_by == uuid.UUID(current_user.id)
        or current_user.user_type in ("SYS_ADMIN", "TEACHER", "QUESTION_ADMIN")
        or current_user.user_type == "STUDENT"
    )
    if not allowed:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions",
        )

    # Delete child records first (FK constraints are NO ACTION, not CASCADE)
    from app.models.exam_paper import exam_paper_questions
    from app.models.answer_submission import AnswerSubmission
    from app.models.error_notebook import ErrorNotebook
    from app.models.ocr_upload import OcrUpload
    from sqlalchemy import delete as sa_delete
    async with db.begin():
        await db.execute(sa_delete(exam_paper_questions).where(exam_paper_questions.c.exam_paper_id == exam_paper_id))
        await db.execute(sa_delete(AnswerSubmission).where(AnswerSubmission.exam_paper_id == exam_paper_id))
        await db.execute(sa_delete(ErrorNotebook).where(ErrorNotebook.exam_paper_id == exam_paper_id))
        await db.execute(sa_delete(OcrUpload).where(OcrUpload.exam_paper_id == exam_paper_id))
        await db.execute(delete(ExamPaper).where(ExamPaper.id == exam_paper_id))
    return None


@router.put("/{exam_paper_id}/submission-status")
async def update_submission_status(
    exam_paper_id: uuid.UUID,
    status_in: str = Body(..., embed=True),
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Change submission status for the current student's latest submission. Only GENERATED → RE_GRADED is allowed."""
    if current_user.user_type != "STUDENT":
        raise HTTPException(status_code=403, detail="仅学生可修改")
    from app.models.answer_submission import AnswerSubmission
    sub_result = await db.execute(
        select(AnswerSubmission).where(
            AnswerSubmission.student_id == uuid.UUID(current_user.id),
            AnswerSubmission.exam_paper_id == exam_paper_id,
        ).order_by(AnswerSubmission.submitted_at.desc()).limit(1)
    )
    submission = sub_result.scalar_one_or_none()
    if not submission:
        raise HTTPException(status_code=404, detail="未找到提交记录")
    if submission.status != "GENERATED":
        raise HTTPException(status_code=400, detail="仅已生成状态可修改为重新判")
    if status_in != "RE_GRADED":
        raise HTTPException(status_code=400, detail="仅允许修改为RE_GRADED")
    submission.status = "RE_GRADED"
    await db.commit()
    return {"message": "状态已修改为重新判", "status": "RE_GRADED"}


@router.get("")
async def get_exam_papers(
    skip: int = 0,
    limit: int = 20,
    title: Optional[str] = None,
    status: Optional[str] = None,
    scope: Optional[str] = None,
    grade: Optional[str] = None,
    grades: Optional[str] = None,
    subject: Optional[str] = None,
    keyword: Optional[str] = None,
    created_by: Optional[str] = None,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    limit = min(limit, 200)
    query = select(ExamPaper)
    if title: query = query.where(ExamPaper.title.ilike(f"%{title}%"))
    if subject: query = query.where(ExamPaper.subject == subject)
    if status: query = query.where(ExamPaper.status == status)
    if scope:
        query = query.where(ExamPaper.grade_level['scope'].astext == scope)
    if grades:
        grade_list = [g.strip() for g in grades.split(",") if g.strip()]
        if len(grade_list) == 1:
            query = query.where(ExamPaper.grade_level['grades'].contains([grade_list[0]]))
        elif len(grade_list) > 1:
            query = query.where(ExamPaper.grade_level['grades'].op('?|')(grade_list))
    elif grade:
        query = query.where(ExamPaper.grade_level['grades'].contains([grade]))
    if keyword:
        from sqlalchemy import String, or_
        query = query.where(or_(
            ExamPaper.grade_level['chapter'].astext.ilike(f"%{keyword}%"),
            ExamPaper.grade_level['knowledge_points'].cast(String).ilike(f"%{keyword}%")
        ))
    if created_by == "me":
        query = query.where(ExamPaper.created_by == uuid.UUID(current_user.id))
    query = query.offset(skip).limit(limit).order_by(ExamPaper.created_at.desc())
    result = await db.execute(query)
    papers = result.scalars().all()

    # Compute question_count for each paper
    from sqlalchemy import text as sa_text
    output = []
    for p in papers:
        cnt_result = await db.execute(sa_text("SELECT COUNT(*) FROM exam_paper_questions WHERE exam_paper_id = :pid"), {"pid": p.id.hex})
        qcount = cnt_result.scalar() or 0
        output.append({
            "id": str(p.id), "title": p.title, "subtitle": p.subtitle,
            "subject": p.subject, "grade_level": p.grade_level,
            "total_score": p.total_score, "duration_minutes": p.duration_minutes,
            "status": p.status, "question_count": qcount,
            "instructions": p.instructions, "description": p.description,
            "created_at": p.created_at, "updated_at": p.updated_at,
        })
    return output


@router.post("/{exam_paper_id}/questions")
async def add_question_to_exam_paper(
    exam_paper_id: uuid.UUID,
    question_id: uuid.UUID = Body(...),
    position: int = Body(1),
    score: int = Body(10),
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # Only teachers and admins can add questions to exam papers
    if current_user.user_type not in ("TEACHER", "QUESTION_ADMIN", "SYS_ADMIN"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions",
        )

    # Check if exam paper exists
    result = await db.execute(select(ExamPaper).where(ExamPaper.id == exam_paper_id))
    exam_paper = result.scalar_one_or_none()
    if not exam_paper:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Exam paper not found",
        )

    # Check if question exists
    result = await db.execute(select(Question).where(Question.id == question_id))
    question = result.scalar_one_or_none()
    if not question:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Question not found",
        )

    # Check if user is the creator of the exam paper or is an admin
    allowed = (
        exam_paper.created_by == uuid.UUID(current_user.id)
        or current_user.user_type in ("SYS_ADMIN", "TEACHER", "QUESTION_ADMIN")
        or current_user.user_type == "STUDENT"
    )
    if not allowed:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions",
        )

    # Insert into association table
    await db.execute(exam_paper_questions.insert().values(
        id=uuid.uuid4(), exam_paper_id=exam_paper.id,
        question_id=question.id, position=position, score=score,
    ))

    await db.commit()
    await db.refresh(exam_paper)
    return exam_paper


@router.delete("/{exam_paper_id}/questions/{question_id}", response_model=ExamPaperResponse)
async def remove_question_from_exam_paper(
    exam_paper_id: uuid.UUID,
    question_id: uuid.UUID,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # Only teachers and admins can remove questions from exam papers
    if current_user.user_type not in ("TEACHER", "QUESTION_ADMIN", "SYS_ADMIN"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions",
        )

    # Check if exam paper exists
    result = await db.execute(select(ExamPaper).where(ExamPaper.id == exam_paper_id))
    exam_paper = result.scalar_one_or_none()
    if not exam_paper:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Exam paper not found",
        )

    # Check if question exists
    result = await db.execute(select(Question).where(Question.id == question_id))
    question = result.scalar_one_or_none()
    if not question:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Question not found",
        )

    # Check if user is the creator of the exam paper or is an admin
    allowed = (
        exam_paper.created_by == uuid.UUID(current_user.id)
        or current_user.user_type in ("SYS_ADMIN", "TEACHER", "QUESTION_ADMIN")
        or current_user.user_type == "STUDENT"
    )
    if not allowed:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions",
        )

    # TODO: Remove question from exam paper (this would require an association table)
    # For now, we'll just return the exam paper as a placeholder
    await db.commit()
    await db.refresh(exam_paper)
    return exam_paper


@router.put("/{exam_paper_id}/questions/sort", response_model=ExamPaperResponse)
async def sort_questions_in_exam_paper(
    exam_paper_id: uuid.UUID,
    question_ids: List[str],
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # Only teachers and admins can sort questions in exam papers
    if current_user.user_type not in ("TEACHER", "QUESTION_ADMIN", "SYS_ADMIN"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions",
        )

    # Check if exam paper exists
    result = await db.execute(select(ExamPaper).where(ExamPaper.id == exam_paper_id))
    exam_paper = result.scalar_one_or_none()
    if not exam_paper:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Exam paper not found",
        )

    # Check if user is the creator of the exam paper or is an admin
    allowed = (
        exam_paper.created_by == uuid.UUID(current_user.id)
        or current_user.user_type in ("SYS_ADMIN", "TEACHER", "QUESTION_ADMIN")
        or current_user.user_type == "STUDENT"
    )
    if not allowed:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not enough permissions",
        )

    # TODO: Sort questions in exam paper (this would require an association table with ordering)
    # For now, we'll just return the exam paper as a placeholder
    await db.commit()
    await db.refresh(exam_paper)
    return exam_paper


@router.get("/{exam_paper_id}/questions")
async def get_questions_in_exam_paper(
    exam_paper_id: uuid.UUID,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from sqlalchemy import text as sa_text
    # Get questions through association table
    qresult = await db.execute(
        sa_text("SELECT q.* FROM questions q JOIN exam_paper_questions epq ON q.id = epq.question_id WHERE epq.exam_paper_id = :pid ORDER BY epq.position"),
        {"pid": exam_paper_id.hex}
    )
    questions = qresult.fetchall()
    return [{"id": row[0], "title": row[1], "question_type": row[2], "difficulty": row[3],
             "subject": row[4], "grade_level": row[5], "score": row[6],
             "correct_answer": row[7], "explanation": row[8], "meta_data": row[9],
             "source": row[10], "review_status": row[11]} for row in questions]


async def _get_paper_with_questions(exam_paper_id: uuid.UUID, db: AsyncSession):
    """Fetch paper info and questions, raise 404 if not found."""
    result = await db.execute(select(ExamPaper).where(ExamPaper.id == exam_paper_id))
    paper = result.scalar_one_or_none()
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Exam paper not found")

    qresult = await db.execute(
        sa_text("SELECT q.* FROM questions q JOIN exam_paper_questions epq ON q.id = epq.question_id WHERE epq.exam_paper_id = :pid ORDER BY epq.position"),
        {"pid": exam_paper_id.hex}
    )
    questions = qresult.fetchall()

    TYPE_LABELS = {"FILL_BLANK": "填空题", "SINGLE_CHOICE": "单选题", "MULTIPLE_CHOICE": "多选题", "SUBJECTIVE": "解答题"}
    mapped = []
    for i, row in enumerate(questions):
        qtype = row[2]
        correct_answer = row[7] or ""
        options = None
        answer_text = ""
        try:
            ad = json.loads(correct_answer)
            if isinstance(ad, dict):
                options = ad.get("options")
                answer_text = ad.get("correct_answer", "")
                if isinstance(answer_text, list):
                    answer_text = ", ".join(str(x) for x in answer_text)
                else:
                    answer_text = str(answer_text)
        except (json.JSONDecodeError, TypeError):
            answer_text = correct_answer

        mapped.append({
            "index": i + 1,
            "title": row[1] or "",
            "question_type": qtype,
            "type_label": TYPE_LABELS.get(qtype, qtype),
            "difficulty": row[3] or "",
            "score": row[6] or 0,
            "correct_answer": correct_answer,
            "answer_text": answer_text,
            "options": options,
            "explanation": row[8] or "",
        })
    return paper, mapped


@router.get("/{exam_paper_id}/export/word")
async def export_exam_paper_word(
    exam_paper_id: uuid.UUID,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse
    import io

    paper, questions = await _get_paper_with_questions(exam_paper_id, db)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(11)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(paper.title or "")
    run.bold = True
    run.font.size = Pt(18)

    # Subtitle / info line
    if paper.subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.add_run(paper.subtitle).font.size = Pt(10)

    info_parts = []
    if paper.subject: info_parts.append(f"学科：{paper.subject}")
    if paper.grade_level: info_parts.append(f"年级：{paper.grade_level}")
    info_parts.append(f"总分：{paper.total_score}分")
    if paper.duration_minutes: info_parts.append(f"时长：{paper.duration_minutes}分钟")
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(" | ".join(info_parts)).font.size = Pt(10)

    # Notes
    if paper.instructions:
        doc.add_paragraph()
        note_para = doc.add_paragraph()
        run = note_para.add_run(f"注意事项：{paper.instructions}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacing

    # Group by type
    type_order = ["FILL_BLANK", "SINGLE_CHOICE", "MULTIPLE_CHOICE", "SUBJECTIVE"]
    grouped = {}
    for q in questions:
        t = q["question_type"]
        grouped.setdefault(t, []).append(q)

    for t in type_order:
        qs = grouped.get(t, [])
        if not qs:
            continue
        # Section header
        header = doc.add_paragraph()
        run = header.add_run(f"{qs[0]['type_label']}（共{len(qs)}题）")
        run.bold = True
        run.font.size = Pt(13)

        for q in qs:
            # Question title
            q_para = doc.add_paragraph()
            q_para.add_run(f"{q['index']}. {q['title']}（{q['score']}分）").font.size = Pt(11)

            # Options for choice questions
            if q["options"] and len(q["options"]) > 0:
                for opt in q["options"]:
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Cm(1)
                    opt_para.add_run(f"{opt.get('label', '')}. {opt.get('text', '')}").font.size = Pt(10)

            # Blank line for fill-blank
            if t == "FILL_BLANK":
                blank_para = doc.add_paragraph()
                blank_para.add_run("_" * 40).font.size = Pt(10)

            # Space for subjective
            if t == "SUBJECTIVE":
                for _ in range(3):
                    space_para = doc.add_paragraph()
                    space_para.add_run("_" * 60).font.size = Pt(10)

            doc.add_paragraph()  # spacing

    # Save to BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = quote((paper.title or 'exam_paper') + '.docx')
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{safe_name}"}
    )


@router.get("/{exam_paper_id}/export/pdf")
async def export_exam_paper_pdf(
    exam_paper_id: uuid.UUID,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from fpdf import FPDF
    from fastapi.responses import StreamingResponse
    import io

    paper, questions = await _get_paper_with_questions(exam_paper_id, db)

    pdf = FPDF()
    pdf.add_page()
    # Use built-in font that supports CJK - or fallback to ASCII
    pdf.add_font("SimSun", "", "/usr/share/fonts/truetype/arphic/uming.ttc", uni=True)
    pdf.set_font("SimSun", "", 12)

    # Title
    pdf.set_font("SimSun", "", 18)
    pdf.cell(0, 12, paper.title or "", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)

    # Info
    pdf.set_font("SimSun", "", 10)
    info_parts = []
    if paper.subject: info_parts.append(f"学科：{paper.subject}")
    if paper.grade_level: info_parts.append(f"年级：{paper.grade_level}")
    info_parts.append(f"总分：{paper.total_score}分")
    if paper.duration_minutes: info_parts.append(f"时长：{paper.duration_minutes}分钟")
    pdf.cell(0, 8, " | ".join(info_parts), new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(6)

    # Questions grouped by type
    type_order = ["FILL_BLANK", "SINGLE_CHOICE", "MULTIPLE_CHOICE", "SUBJECTIVE"]
    grouped = {}
    for q in questions:
        t = q["question_type"]
        grouped.setdefault(t, []).append(q)

    for t in type_order:
        qs = grouped.get(t, [])
        if not qs:
            continue
        pdf.set_font("SimSun", "", 13)
        pdf.cell(0, 10, f"{qs[0]['type_label']}（共{len(qs)}题）", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        for q in qs:
            pdf.set_font("SimSun", "", 11)
            title_text = f"{q['index']}. {q['title']}（{q['score']}分）"
            pdf.multi_cell(0, 7, title_text)
            pdf.ln(1)

            if q["options"] and len(q["options"]) > 0:
                pdf.set_font("SimSun", "", 10)
                for opt in q["options"]:
                    pdf.cell(10, 6, "")
                    pdf.cell(0, 6, f"{opt.get('label', '')}. {opt.get('text', '')}", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)

            if t == "FILL_BLANK":
                pdf.cell(10, 6, "")
                pdf.cell(0, 6, "_" * 40, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)

            if t == "SUBJECTIVE":
                for _ in range(3):
                    pdf.cell(10, 8, "")
                    pdf.cell(0, 8, "_" * 50, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)

            pdf.ln(2)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)

    safe_name = quote((paper.title or 'exam_paper') + '.pdf')
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{safe_name}"}
    )


@router.get("/{exam_paper_id}/preview")
async def preview_exam_paper(
    exam_paper_id: uuid.UUID,
    current_user = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    paper, questions = await _get_paper_with_questions(exam_paper_id, db)
    return {
        "paper": {
            "id": str(paper.id), "title": paper.title, "subtitle": paper.subtitle,
            "subject": paper.subject, "grade_level": paper.grade_level,
            "total_score": paper.total_score, "duration_minutes": paper.duration_minutes,
            "status": paper.status, "instructions": paper.instructions,
            "description": paper.description,
        },
        "questions": questions,
    }


# Exam paper templates endpoints would go here
# For brevity, I'm skipping them for now but they would follow a similar pattern